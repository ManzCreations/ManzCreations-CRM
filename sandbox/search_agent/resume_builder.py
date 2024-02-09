import os
import random
import subprocess

from docx import Document
from faker import Faker
from tqdm import tqdm

# Initialize Faker
fake = Faker()


def generate_resume_content():
    print("Generating resume content...")
    # Basic Info
    name = fake.name()
    address = fake.address().replace("\\n", ", ")
    email = fake.email()
    phone = fake.phone_number()

    # Education
    universities = ["University of California", "Massachusetts Institute of Technology", "Stanford University",
                    "Harvard University", "Princeton University"]
    degrees = ["B.S.", "M.S.", "Ph.D."]
    majors = ["Computer Science", "Mechanical Engineering", "Electrical Engineering", "Business Administration",
              "Chemistry"]
    education = [{
        "university": random.choice(universities),
        "degree": random.choice(degrees),
        "major": random.choice(majors),
        "graduation_month": random.choice(["June", "July", "December"]),
        "graduation_year": random.randint(2010, 2023)
    } for _ in range(2)]  # Generating two degrees

    # Experience
    companies = ["Google", "Amazon", "Facebook", "Microsoft", "Tesla"]
    titles = ["Software Engineer", "Data Analyst", "Project Manager", "Product Designer", "Research Scientist"]
    experience = [{
        "company": random.choice(companies),
        "title": random.choice(titles),
        "start_month": random.choice(["January", "February", "March"]),
        "start_year": random.randint(2015, 2022),
        "end_month": random.choice(["April", "May", "June"]),
        "end_year": random.randint(2023, 2024),
        "description": fake.text(max_nb_chars=200)
    } for _ in range(3)]  # Generating three job experiences

    # Projects
    project_names = ["Web Development Platform", "Machine Learning Model for Stock Prediction",
                     "IoT Home Security System", "Mobile App for Health Monitoring", "Blockchain-Based Voting System"]
    projects = [{
        "name": random.choice(project_names),
        "tools": ["Python", "Java", "React", "Node.js", "TensorFlow", "Solidity"],
        "description": fake.text(max_nb_chars=150),
        "duration": f"{random.choice(['January', 'February', 'March'])} {random.randint(2021, 2023)} - "
                    f"{random.choice(['April', 'May', 'June'])} {random.randint(2023, 2024)}"
    } for _ in range(2)]  # Generating two projects

    # Compile resume
    resume = {
        "name": name,
        "address": address,
        "email": email,
        "phone": phone,
        "education": education,
        "experience": experience,
        "projects": projects
    }

    return resume


def generate_docx(resume, filename):
    print(f"Generating DOCX file: {filename}...")
    try:
        doc = Document()

        doc.add_heading(resume['name'], 0)
        doc.add_paragraph(f"Address: {resume['address']}")
        doc.add_paragraph(f"Email: {resume['email']}")
        doc.add_paragraph(f"Phone: {resume['phone']}")

        doc.add_heading('Education', level=1)
        for edu in resume['education']:
            doc.add_paragraph(
                f"{edu['degree']} in {edu['major']}, {edu['university']} - "
                f"{edu['graduation_month']} {edu['graduation_year']}",
                style='ListBullet')

        doc.add_heading('Experience', level=1)
        for exp in resume['experience']:
            doc.add_paragraph(
                f"{exp['title']} at {exp['company']}, {exp['start_month']} {exp['start_year']} - "
                f"{exp['end_month']} {exp['end_year']}",
                style='ListBullet')
            doc.add_paragraph(exp['description'], style='ListBullet')

        doc.add_heading('Projects', level=1)
        for proj in resume['projects']:
            doc.add_paragraph(f"{proj['name']}, {proj['duration']}", style='ListBullet')
            doc.add_paragraph(f"Tools: {', '.join(proj['tools'])}", style='ListBullet')
            doc.add_paragraph(proj['description'], style='ListBullet')

        doc.save(filename)
        print(f".DOCX generated successfully: {filename}")
    except Exception as e:
        print(f"Failed to generate .DOCX: {filename}. Error: {e}.")


def convert_docx_to_pdf(docx_filename, pdf_filename):
    print(f"Converting {docx_filename} to PDF: {pdf_filename}...")
    # Assuming using a command-line tool like LibreOffice or Pandoc for conversion
    # Example: LibreOffice --convert-to pdf:writer_pdf_Export --outdir <output_dir> <docx_filename>
    try:
        subprocess.run(["soffice", "--convert-to", "pdf", "--outdir",
                        os.path.dirname(pdf_filename), docx_filename], check=True)
        print(f"PDF conversion successful: {pdf_filename}")
    except subprocess.CalledProcessError as e:
        print(f"Error converting {docx_filename} to PDF: {e}")


def main(output_dir="resumes"):
    print(f"Creating output directory: {output_dir}...")

    os.makedirs(output_dir, exist_ok=True)

    resume_count = 20

    # Generate 20 DOCX resumes, convert them to PDF, then delete the DOCX file
    print("Starting to generate PDF resumes from DOCX files...")
    for _ in tqdm(range(resume_count), desc="Generating .PDF resumes"):
        resume_content = generate_resume_content()
        last_name = resume_content['name'].split()[-1]
        try:
            docx_filename = f"{output_dir}/{last_name}_Resume.docx"
            pdf_filename = f"{output_dir}/{last_name}_Resume.pdf"

            generate_docx(resume_content, docx_filename)
            convert_docx_to_pdf(docx_filename, pdf_filename)

            # Delete the DOCX file after conversion
            print(f"Deleting the DOCX file: {docx_filename}...")
            os.remove(docx_filename)
        except Exception as e:
            print(f"An error occurred during .pdf resume generation for {last_name}: {e}")

    # Generate 20 DOCX resumes
    print("Starting to generate DOCX resumes...")
    for _ in tqdm(range(resume_count), desc="Generating .DOCX resumes"):
        resume_content = generate_resume_content()
        last_name = resume_content['name'].split()[-1]
        try:
            docx_filename = f"{output_dir}/{last_name}_Resume.docx"
            generate_docx(resume_content, docx_filename)
        except Exception as e:
            print(f"An error occurred during .docx resume generation for {last_name}: {e}")

    print("Resume generation process completed.")


if __name__ == "__main__":
    main()
