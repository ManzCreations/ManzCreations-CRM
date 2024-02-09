import os
import random

from docx import Document
from faker import Faker
from tqdm import tqdm

# Initialize Faker
fake = Faker()


def generate_job_description_content():
    print("Generating job description content...")
    job_roles = {
        "Electrician": {
            "Responsibilities": [
                "Install, maintain, and repair electrical control, wiring, and lighting systems.",
                "Read technical diagrams and blueprints.",
                "Perform general electrical maintenance.",
                "Inspect transformers and circuit breakers.",
                "Troubleshoot electrical issues using appropriate testing devices.",
                "Follow National Electrical Code state and local building regulations.",
                "Execute plans of electrical wiring for well-functioning lighting and other electrical systems.",
                "Connect wiring in electrical circuits and networks ensuring compatibility."
            ],
            "Requirements": [
                "Completion of electrician apprenticeship.",
                "Valid electrician license.",
                "Proven experience as an electrician.",
                "In-depth understanding of safety procedures and legal regulations.",
                "Excellent critical thinking and problem-solving ability.",
                "Ability to work in a team and to communicate well."
            ]
        },
        "Software Developer": {
            "Responsibilities": [
                "Develop and implement software systems and applications.",
                "Analyze user needs and software requirements.",
                "Test and maintain software products for optimization.",
                "Recommend improvements to existing software.",
                "Collaborate with other developers, UX designers, business and systems analysts.",
                "Write well-designed, testable code."
            ],
            "Requirements": [
                "Bachelor's degree in Computer Science or related field.",
                "Proficiency in coding languages such as Java, Python, and C#.",
                "Experience with software design and development in a test-driven environment.",
                "Knowledge of coding languages and frameworks/systems (e.g., AngularJS, Git).",
                "Excellent communication skills."
            ]
        },
        "Marketing Manager": {
            "Responsibilities": [
                "Develop marketing strategies to increase company revenue.",
                "Oversee social media marketing campaigns.",
                "Analyze market trends and competitors.",
                "Manage the marketing team to achieve targets.",
                "Organize and oversee advertising/communication campaigns.",
                "Conduct market research and analysis to evaluate trends."
            ],
            "Requirements": [
                "Bachelor's degree in Marketing or related field.",
                "Proven experience in marketing management.",
                "Strong analytical skills and data-driven thinking.",
                "Up-to-date with the latest trends and best practices in online marketing.",
                "Excellent communication skills."
            ]
        },
        "Human Resources Manager": {
            "Responsibilities": [
                "Develop HR strategies aligned with the business strategy.",
                "Manage employee relations and grievances.",
                "Oversee recruitment and selection process.",
                "Manage performance appraisal system.",
                "Assess training needs and coordinate learning and development initiatives.",
                "Ensure legal compliance throughout human resource management."
            ],
            "Requirements": [
                "Proven experience as an HR Manager.",
                "Knowledge of HR systems and databases.",
                "Ability to architect strategy along with leadership skills.",
                "Excellent active listening, negotiation, and presentation skills.",
                "Competence to build and effectively manage interpersonal relationships."
            ]
        },
        # Additional roles can be added here
    }

    # Randomly select a job role
    role = random.choice(list(job_roles.keys()))
    responsibilities = random.sample(job_roles[role]["Responsibilities"],
                                     k=random.randint(3, len(job_roles[role]["Responsibilities"])))
    requirements = random.sample(job_roles[role]["Requirements"],
                                 k=random.randint(3, len(job_roles[role]["Requirements"])))

    # Compile job description
    job_description = {
        "role": role,
        "responsibilities": responsibilities,
        "requirements": requirements
    }

    return job_description


def generate_docx(job_description, filename):
    print(f"Generating DOCX file for job description: {filename}...")
    doc = Document()

    doc.add_heading(f"{job_description['role']} Job Description", 0)

    doc.add_heading('Responsibilities:', level=1)
    for responsibility in job_description['responsibilities']:
        doc.add_paragraph(responsibility, style='ListBullet')

    doc.add_heading('Requirements:', level=1)
    for requirement in job_description['requirements']:
        doc.add_paragraph(requirement, style='ListBullet')

    doc.save(filename)
    print(f".DOCX generated successfully for job description: {filename}")


def main(output_dir="job_descriptions"):
    print(f"Creating output directory: {output_dir}...")

    # Ensure the output directory exists
    os.makedirs(output_dir, exist_ok=True)

    job_description_count = 5

    # Generate 5 DOCX job descriptions
    print("Starting to generate DOCX job descriptions...")
    file_counter = 1
    for _ in tqdm(range(job_description_count), desc="Generating .DOCX job descriptions"):
        job_description_content = generate_job_description_content()
        if job_description_content:
            role = job_description_content['role'].replace(" ", "_")  # Replace spaces with underscores for filename
            filename = f"{output_dir}/{role}_{file_counter}_Job_Description.docx"

            generate_docx(job_description_content, filename)
            file_counter += 1
        else:
            print("Failed to generate job description content, skipping DOCX generation.")

    print("Job description generation process completed.")


if __name__ == "__main__":
    main()
